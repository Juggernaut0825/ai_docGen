#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
主程序：AI文档生成器 - 增强版
支持位置感知的AI字段映射和智能模板填充
"""

import os
import json
import logging
import subprocess
from datetime import datetime
from typing import Dict, Any, List
from docx import Document
from openai import OpenAI

# 导入新的提示词工具
from prompt_utils import PromptTemplates, PromptHelper

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    datefmt='%H:%M:%S'
)
logger = logging.getLogger(__name__)

class EnhancedAIDocGenerator:
    """增强版AI文档生成器 - 支持位置感知映射"""
    
    def __init__(self, api_key: str):
        """初始化OpenRouter客户端"""
        self.client = OpenAI(
            base_url="https://openrouter.ai/api/v1",
            api_key=api_key,
        )
        self.model = "google/gemini-2.5-pro-preview"
        self.prompt_templates = PromptTemplates()
        self.prompt_helper = PromptHelper()
        logger.info("🤖 增强版AI生成器初始化完成")
    
    def convert_doc_to_docx(self, doc_path: str) -> str:
        """
        使用LibreOffice将.doc文件转换为.docx文件
        
        Args:
            doc_path: .doc文件路径
            
        Returns:
            转换后的.docx文件路径
        """
        logger.info("🔄 开始DOC到DOCX转换...")
        
        if not os.path.exists(doc_path):
            logger.error(f"❌ DOC文件不存在: {doc_path}")
            raise FileNotFoundError(f"DOC文件不存在: {doc_path}")
        
        # 生成输出文件名
        docx_path = doc_path.replace('.doc', '_converted.docx')
        
        try:
            # 检查LibreOffice是否可用
            logger.info("🔍 检查LibreOffice可用性...")
            
            # 尝试多个可能的LibreOffice路径
            libreoffice_paths = [
                '/Applications/LibreOffice.app/Contents/MacOS/soffice',  # macOS
                'libreoffice',  # Linux/Windows PATH
                'soffice',  # 备用命令
            ]
            
            libreoffice_cmd = None
            for path in libreoffice_paths:
                try:
                    result = subprocess.run([path, '--version'], 
                                          capture_output=True, 
                                          text=True, 
                                          timeout=10)
                    if result.returncode == 0:
                        libreoffice_cmd = path
                        logger.info(f"✅ 找到LibreOffice: {path}")
                        break
                except (FileNotFoundError, subprocess.TimeoutExpired):
                    continue
            
            if not libreoffice_cmd:
                logger.error("❌ 未找到LibreOffice，请确保已安装LibreOffice")
                raise RuntimeError("LibreOffice未安装或不可用")
            
            # 执行转换
            logger.info(f"📄 正在转换: {doc_path} -> {docx_path}")
            
            # 删除已存在的输出文件
            if os.path.exists(docx_path):
                os.remove(docx_path)
                logger.info("🗑️ 删除已存在的转换文件")
            
            # LibreOffice转换命令
            cmd = [
                libreoffice_cmd,
                '--headless',
                '--convert-to', 'docx',
                '--outdir', os.path.dirname(doc_path),
                doc_path
            ]
            
            logger.info(f"🔧 执行命令: {' '.join(cmd)}")
            
            result = subprocess.run(cmd, 
                                  capture_output=True, 
                                  text=True, 
                                  timeout=30)
            
            if result.returncode != 0:
                logger.error(f"❌ LibreOffice转换失败: {result.stderr}")
                raise RuntimeError(f"LibreOffice转换失败: {result.stderr}")
            
            # 检查转换后的文件
            expected_docx = doc_path.replace('.doc', '.docx')
            if os.path.exists(expected_docx):
                # 重命名为我们期望的文件名
                if expected_docx != docx_path:
                    os.rename(expected_docx, docx_path)
                
                logger.info(f"✅ 转换成功: {docx_path}")
                return docx_path
            else:
                logger.error(f"❌ 转换后的文件未找到: {expected_docx}")
                raise RuntimeError("转换后的文件未找到")
                
        except subprocess.TimeoutExpired:
            logger.error("❌ LibreOffice转换超时")
            raise RuntimeError("LibreOffice转换超时")
        except Exception as e:
            logger.error(f"❌ 转换过程中出错: {e}")
            raise
    
    def stage1_analyze_template_with_position(self, template_path: str) -> Dict[str, str]:
        """
        阶段1：增强版模板分析 - 提取位置信息和上下文
        """
        logger.info("🔍 阶段1：开始位置感知的模板结构分析...")
        
        try:
            # 读取Word文档内容
            doc = Document(template_path)
            template_content = ""
            
            logger.info(f"📄 正在读取模板文件: {template_path}")
            
            # 增强版模板内容提取 - 包含更多上下文信息
            table_count = 0
            for table in doc.tables:
                table_count += 1
                logger.info(f"📋 处理第 {table_count} 个表格")
                template_content += f"\n=== 表格 {table_count} ===\n"
                
                for row_idx, row in enumerate(table.rows):
                    row_content = ""
                    for cell_idx, cell in enumerate(row.cells):
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_content += f"[Row{row_idx+1}Col{cell_idx+1}]: {cell_text} | "
                    
                    if row_content:
                        template_content += f"第{row_idx+1}行: {row_content}\n"
                        
                        # 添加上下文分析
                        if row_idx > 0:
                            prev_row_content = ""
                            for cell_idx, cell in enumerate(table.rows[row_idx-1].cells):
                                prev_cell_text = cell.text.strip()
                                if prev_cell_text:
                                    prev_row_content += f"{prev_cell_text} | "
                            if prev_row_content:
                                template_content += f"  上方行内容: {prev_row_content}\n"
            
            logger.info(f"📊 增强版模板内容提取完成，共 {table_count} 个表格")
            
            # 使用增强版提示词分析模板结构
            prompt = self.prompt_templates.get_template_analysis_prompt(template_content)
            
            logger.info("🧠 正在调用AI进行位置感知的模板字段分析...")
            
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[{"role": "user", "content": prompt}],
                extra_headers={
                    "HTTP-Referer": "ai-doc-generator",
                    "X-Title": "AI Document Generator Enhanced",
                }
            )
            
            # 解析返回的JSON
            json_text = self.prompt_helper.extract_json_from_response(response.choices[0].message.content)
            
            # 验证JSON有效性
            if not self.prompt_helper.validate_json_structure(json_text):
                logger.warning("⚠️ AI返回的JSON格式无效，使用降级结构")
                template_structure = self.prompt_helper.create_fallback_structure(template_content)
            else:
                template_structure = json.loads(json_text)
            
            logger.info(f"✅ 成功提取 {len(template_structure)} 个位置感知字段:")
            for key, value in template_structure.items():
                logger.info(f"   📌 {key}: {value}")
            
            # 记录字段统计信息
            self._log_field_statistics(template_structure)
            
            return template_structure
            
        except Exception as e:
            logger.error(f"❌ 阶段1错误: {e}")
            # 返回降级结构
            fallback_structure = self.prompt_helper.create_fallback_structure("")
            logger.warning("⚠️ 使用降级模板结构")
            return fallback_structure
    
    def stage2_load_json_data(self, json_file_path: str) -> Dict[str, str]:
        """
        阶段2：从JSON文件加载数据（增强版日志）
        """
        logger.info("📂 阶段2：开始加载JSON数据...")
        
        try:
            if not os.path.exists(json_file_path):
                logger.error(f"❌ JSON文件不存在: {json_file_path}")
                return {}
            
            logger.info(f"📄 正在读取JSON文件: {json_file_path}")
            
            with open(json_file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            logger.info(f"✅ 成功加载 {len(data)} 个数据字段:")
            for key, value in data.items():
                preview = value[:50] + "..." if len(str(value)) > 50 else str(value)
                logger.info(f"   📌 {key}: {preview}")
            
            # 记录数据字段统计
            self._log_data_statistics(data)
            
            return data
            
        except Exception as e:
            logger.error(f"❌ 阶段2错误: {e}")
            return {}
    
    def stage2_5_enhanced_ai_field_mapping(self, template_structure: Dict[str, str], input_data: Dict[str, str]) -> Dict[str, str]:
        """
        阶段2.5：增强版AI智能字段映射
        
        Args:
            template_structure: 位置感知的模板字段结构
            input_data: 输入数据
            
        Returns:
            映射后的数据，使用位置感知的模板字段名作为键
        """
        logger.info("🧠 阶段2.5：开始增强版AI字段映射...")
        
        try:
            # 构建增强版AI映射提示
            base_prompt = self.prompt_templates.get_field_mapping_prompt(template_structure, input_data)
            enhanced_prompt = self.prompt_templates.enhance_mapping_prompt_with_examples(base_prompt)
            
            logger.info("🧠 正在调用AI进行增强版字段映射...")
            logger.info(f"📊 模板字段数量: {len(template_structure)}")
            logger.info(f"📊 输入数据字段数量: {len(input_data)}")
            
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[{"role": "user", "content": enhanced_prompt}],
                extra_headers={
                    "HTTP-Referer": "ai-doc-generator",
                    "X-Title": "AI Document Generator Enhanced",
                }
            )
            
            # 解析返回的JSON
            json_text = self.prompt_helper.extract_json_from_response(response.choices[0].message.content)
            
            # 验证JSON有效性
            if not self.prompt_helper.validate_json_structure(json_text):
                logger.warning("⚠️ AI映射返回的JSON格式无效，尝试直接映射")
                mapped_data = self._fallback_field_mapping(template_structure, input_data)
            else:
                mapped_data = json.loads(json_text)
            
            logger.info(f"✅ 成功映射 {len(mapped_data)} 个字段:")
            for key, value in mapped_data.items():
                preview = value[:50] + "..." if len(str(value)) > 50 else str(value)
                logger.info(f"   🔗 {key}: {preview}")
            
            # 详细的映射统计和验证
            self._log_mapping_statistics(template_structure, input_data, mapped_data)
            
            return mapped_data
            
        except Exception as e:
            logger.error(f"❌ 阶段2.5错误: {e}")
            logger.warning("⚠️ AI字段映射失败，使用降级映射策略")
            return self._fallback_field_mapping(template_structure, input_data)
    
    def stage3_position_aware_template_filling(self, template_path: str, output_path: str, mapped_data: Dict[str, str], template_structure: Dict[str, str]):
        """
        阶段3：位置感知的智能模板填充
        
        Args:
            template_path: 模板文件路径
            output_path: 输出文件路径
            mapped_data: 位置感知映射后的数据
            template_structure: 模板结构（用于验证）
        """
        logger.info("📝 阶段3：开始位置感知的智能模板填充...")
        
        if not os.path.exists(template_path):
            logger.error(f"❌ 模板文件未找到: {template_path}")
            return False

        try:
            logger.info(f"📄 正在打开模板: {template_path}")
            doc = Document(template_path)

            if not doc.tables:
                logger.error("❌ 文档中未找到任何表格")
                return False

            table = doc.tables[0]
            filled_fields = []
            skipped_fields = []
            position_matches = {}

            logger.info("🔍 开始位置感知的智能搜索和填充...")

            # 构建位置映射表
            for key in mapped_data.keys():
                if mapped_data[key]:  # 只处理有值的字段
                    position_info = self._parse_position_key(key)
                    if position_info:
                        position_matches[key] = position_info
                        logger.info(f"🎯 位置解析: {key} -> Row{position_info['row']}, Col{position_info['col']}, Context: {position_info.get('context', 'N/A')}")

            # 遍历表格进行位置匹配填充
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    
                    # 查找匹配的位置字段
                    for field_key, field_value in mapped_data.items():
                        if not field_value:  # 跳过空值
                            continue
                            
                        position_info = position_matches.get(field_key)
                        if not position_info:
                            continue
                        
                        # 位置匹配逻辑
                        if self._is_position_match(row_idx, cell_idx, cell_text, position_info, row):
                            try:
                                success = self._fill_cell_by_position(cell, row, cell_idx, field_value, position_info)
                                if success:
                                    filled_fields.append(f"{field_key} -> {field_value[:50]}{'...' if len(field_value) > 50 else ''}")
                                    logger.info(f"   ✏️ 位置填充成功: {field_key}")
                                else:
                                    skipped_fields.append(f"{field_key}: 填充失败")
                                    logger.warning(f"   ⚠️ 位置填充失败: {field_key}")
                            except Exception as e:
                                skipped_fields.append(f"{field_key}: {str(e)}")
                                logger.error(f"   ❌ 填充异常: {field_key} - {e}")

            # 保存文档
            doc.save(output_path)
            
            # 详细的填充结果统计
            logger.info(f"✅ 文档已成功生成: {output_path}")
            logger.info(f"📊 共填充 {len(filled_fields)} 个字段:")
            for field in filled_fields:
                logger.info(f"   ✓ {field}")
            
            if skipped_fields:
                logger.warning(f"⚠️ 跳过 {len(skipped_fields)} 个字段:")
                for field in skipped_fields:
                    logger.warning(f"   ⏭️ {field}")
            
            # 验证未填充的模板字段
            self._validate_unfilled_fields(template_structure, filled_fields)
            
            return True

        except Exception as e:
            logger.error(f"❌ 阶段3错误: {e}")
            return False

    def _parse_position_key(self, position_key: str) -> Dict[str, Any]:
        """
        解析位置键值，提取行列和上下文信息
        
        Args:
            position_key: 如 "row_3_col_2_编号" 或 "row_4_left_原形制_现场复核情况"
            
        Returns:
            解析后的位置信息字典
        """
        try:
            parts = position_key.split('_')
            if len(parts) < 4:
                return None
            
            row = int(parts[1]) - 1  # 转换为0索引
            
            # 处理不同的位置格式
            if parts[2] == "col":
                col = int(parts[3]) - 1  # 转换为0索引
                field_name = '_'.join(parts[4:])
                return {
                    'row': row,
                    'col': col,
                    'field_name': field_name,
                    'fill_type': 'next_cell'
                }
            elif parts[2] == "left":
                context = parts[3]
                field_name = '_'.join(parts[4:])
                return {
                    'row': row,
                    'col': -1,  # 表示需要根据上下文查找
                    'context': context,
                    'field_name': field_name,
                    'fill_type': 'same_cell_with_context'
                }
            
            return None
        except (ValueError, IndexError):
            return None

    def _is_position_match(self, row_idx: int, cell_idx: int, cell_text: str, position_info: Dict[str, Any], row) -> bool:
        """
        判断当前位置是否匹配目标填充字段
        
        Args:
            row_idx: 当前行索引
            cell_idx: 当前列索引
            cell_text: 当前单元格文本
            position_info: 位置信息
            row: 当前行对象
            
        Returns:
            是否匹配
        """
        if position_info['fill_type'] == 'next_cell':
            # 检查标签字段位置匹配
            return (row_idx == position_info['row'] and 
                    cell_idx == position_info['col'] and 
                    position_info['field_name'] in cell_text)
        
        elif position_info['fill_type'] == 'same_cell_with_context':
            # 检查上下文相关的填充
            if row_idx == position_info['row'] and position_info['field_name'] in cell_text:
                # 检查左侧单元格是否包含上下文
                if cell_idx > 0:
                    left_cell_text = row.cells[cell_idx-1].text
                    return position_info['context'] in left_cell_text
                    
        return False

    def _fill_cell_by_position(self, cell, row, cell_idx: int, field_value: str, position_info: Dict[str, Any]) -> bool:
        """
        根据位置信息填充单元格
        
        Args:
            cell: 当前单元格对象
            row: 当前行对象
            cell_idx: 单元格索引
            field_value: 要填充的值
            position_info: 位置信息
            
        Returns:
            是否填充成功
        """
        try:
            if position_info['fill_type'] == 'next_cell':
                # 填充下一个单元格
                if cell_idx + 1 < len(row.cells):
                    row.cells[cell_idx + 1].text = field_value
                    return True
            
            elif position_info['fill_type'] == 'same_cell_with_context':
                # 在同一单元格中添加内容
                cell.add_paragraph(field_value)
                return True
                
        except Exception as e:
            logger.error(f"❌ 填充单元格时出错: {e}")
            
        return False

    def _fallback_field_mapping(self, template_structure: Dict[str, str], input_data: Dict[str, str]) -> Dict[str, str]:
        """
        降级字段映射策略（当AI映射失败时使用）
        """
        logger.info("🔄 执行降级字段映射策略...")
        
        mapped_data = {}
        
        # 基础映射规则
        mapping_rules = {
            "serial_number": ["编号", "number", "id"],
            "project_name": ["项目名称", "name", "title"],  
            "review_date": ["复核日期", "date", "check_date"],
            "original_condition_review": ["原形制", "original_state", "original_form"],
            "damage_assessment_review": ["病害和残损", "damage", "deterioration"],
            "repair_plan_review": ["修缮做法", "repair_method", "repair_plan"],
            "project_lead": ["项目负责人", "project_manager", "manager", "lead"],
            "reviewer": ["复核人员", "reviewers", "checker"]
        }
        
        # 反向映射：从输入数据字段到模板字段
        for template_key in template_structure.keys():
            for input_key, input_value in input_data.items():
                if input_value:  # 只处理有值的字段
                    # 直接匹配
                    if input_key in template_key or any(rule in template_key for rule_list in mapping_rules.values() for rule in rule_list):
                        mapped_data[template_key] = input_value
                        break
                    
                    # 语义匹配
                    for semantic_key, possible_names in mapping_rules.items():
                        if input_key in possible_names or input_key == semantic_key:
                            if any(name in template_key for name in possible_names):
                                mapped_data[template_key] = input_value
                                break
        
        logger.info(f"🔄 降级映射完成，映射了 {len(mapped_data)} 个字段")
        return mapped_data

    def _log_field_statistics(self, template_structure: Dict[str, str]):
        """记录字段统计信息"""
        field_types = {'basic_info': 0, 'review_situation': 0, 'personnel': 0, 'other': 0}
        
        for key in template_structure.keys():
            if any(info in key for info in ['编号', '项目名称', '复核日期']):
                field_types['basic_info'] += 1
            elif '现场复核情况' in key:
                field_types['review_situation'] += 1
            elif any(person in key for person in ['负责人', '复核人员']):
                field_types['personnel'] += 1
            else:
                field_types['other'] += 1
        
        logger.info(f"📊 字段类型统计: {field_types}")

    def _log_data_statistics(self, input_data: Dict[str, str]):
        """记录输入数据统计信息"""
        total_chars = sum(len(str(value)) for value in input_data.values())
        non_empty_fields = sum(1 for value in input_data.values() if value)
        
        logger.info(f"📊 输入数据统计: 总字段 {len(input_data)}, 非空字段 {non_empty_fields}, 总字符数 {total_chars}")

    def _log_mapping_statistics(self, template_structure: Dict[str, str], input_data: Dict[str, str], mapped_data: Dict[str, str]):
        """记录映射统计信息"""
        mapped_count = len(mapped_data)
        template_count = len(template_structure)
        input_count = len(input_data)
        
        mapped_with_values = sum(1 for value in mapped_data.values() if value)
        
        logger.info(f"📊 映射统计: 模板字段 {template_count}, 输入字段 {input_count}, 映射字段 {mapped_count}, 有值字段 {mapped_with_values}")
        
        # 检查未映射的模板字段
        unmapped_template = [key for key in template_structure.keys() if key not in mapped_data or not mapped_data[key]]
        if unmapped_template:
            logger.warning(f"⚠️ 未映射的模板字段 ({len(unmapped_template)}): {unmapped_template}")
        
        # 检查未使用的输入字段
        used_input_values = set(mapped_data.values())
        unused_input = [key for key, value in input_data.items() if value and value not in used_input_values]
        if unused_input:
            logger.warning(f"⚠️ 未使用的输入字段 ({len(unused_input)}): {unused_input}")

    def _validate_unfilled_fields(self, template_structure: Dict[str, str], filled_fields: List[str]):
        """验证未填充的字段"""
        filled_field_keys = [field.split(' -> ')[0] for field in filled_fields]
        unfilled_keys = [key for key in template_structure.keys() if key not in filled_field_keys]
        
        if unfilled_keys:
            logger.warning(f"⚠️ 模板中有 {len(unfilled_keys)} 个字段未被填充:")
            for key in unfilled_keys:
                logger.warning(f"   🔍 未填充: {key}")

    def run_enhanced_workflow(self, doc_template_path: str, json_input_path: str, output_path: str):
        """
        运行增强版的完整工作流程
        """
        logger.info("🚀 开始增强版AI文档生成流程")
        logger.info("=" * 60)
        
        start_time = datetime.now()
        
        try:
            # 阶段0：DOC转DOCX转换
            docx_template_path = self.convert_doc_to_docx(doc_template_path)
            logger.info("=" * 30)
            
            # 阶段1：位置感知的模板分析
            template_structure = self.stage1_analyze_template_with_position(docx_template_path)
            logger.info("=" * 30)
            
            # 阶段2：加载JSON数据
            input_data = self.stage2_load_json_data(json_input_path)
            logger.info("=" * 30)
            
            # 阶段2.5：增强版AI字段映射
            mapped_data = self.stage2_5_enhanced_ai_field_mapping(template_structure, input_data)
            logger.info("=" * 30)
            
            # 阶段3：位置感知的模板填充
            success = self.stage3_position_aware_template_filling(docx_template_path, output_path, mapped_data, template_structure)
            
            end_time = datetime.now()
            duration = (end_time - start_time).total_seconds()
            
            logger.info("=" * 60)
            if success:
                logger.info("🎉 增强版AI文档生成流程完成!")
                logger.info(f"⏱️ 总用时: {duration:.2f} 秒")
                logger.info(f"📄 输出文件: {output_path}")
                logger.info(f"🔄 中间转换文件: {docx_template_path}")
            else:
                logger.error("❌ 文档生成失败")
            
            return success
            
        except Exception as e:
            logger.error(f"❌ 增强版工作流程失败: {e}")
            return False


def main():
    """主函数"""
    print("🚀 增强版AI文档生成器 - 主程序")
    print("=" * 50)
    
    # 配置
    API_KEY = "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"
    
    # 文件路径
    doc_template_path = "template_test.doc"
    json_input_path = "sample_input.json"
    output_path = f"增强版AI生成文档_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    
    # 检查文件
    if not os.path.exists(doc_template_path):
        logger.error(f"❌ DOC模板文件不存在: {doc_template_path}")
        return
    
    if not os.path.exists(json_input_path):
        logger.error(f"❌ JSON输入文件不存在: {json_input_path}")
        return
    
    # 初始化增强版生成器
    try:
        generator = EnhancedAIDocGenerator(API_KEY)
    except Exception as e:
        logger.error(f"❌ 增强版生成器初始化失败: {e}")
        return
    
    # 运行增强版完整流程
    success = generator.run_enhanced_workflow(
        doc_template_path=doc_template_path,
        json_input_path=json_input_path,
        output_path=output_path
    )
    
    if success:
        print(f"\n✅ 成功！生成的文档: {output_path}")
    else:
        print("\n❌ 失败！请检查日志信息")


if __name__ == "__main__":
    main() 